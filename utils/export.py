"""
Export Utility — PDF and Word (.docx) export for case files.

Generates:
  - Evidence checklist (PDF + DOCX)
  - Written submissions (PDF + DOCX)
  - Full case file (combined PDF)
  - Paper book index (DOCX)

Usage:
    from utils.export import build_evidence_pdf, build_submission_docx
    pdf_bytes = build_evidence_pdf(case, evidence_items)
    docx_bytes = build_submission_docx(case, submission_text)
"""

from __future__ import annotations
import io
import os
import sys
from datetime import datetime

sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))


# ─────────────────────────────────────────────────────────────────────────────
# PDF — Evidence Checklist
# ─────────────────────────────────────────────────────────────────────────────

def build_evidence_pdf(case: dict, evidence_items: list[dict]) -> bytes:
    """
    Build a formatted PDF evidence checklist.
    Returns raw PDF bytes suitable for st.download_button.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    Table, TableStyle, HRFlowable)
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    buf    = io.BytesIO()
    doc    = SimpleDocTemplate(buf, pagesize=A4,
                                topMargin=2*cm, bottomMargin=2*cm,
                                leftMargin=2.5*cm, rightMargin=2.5*cm)
    styles = getSampleStyleSheet()

    # Custom styles
    title_style = ParagraphStyle("Title2", parent=styles["Title"],
                                  fontSize=16, spaceAfter=6,
                                  textColor=colors.HexColor("#1E3A5F"))
    heading_style = ParagraphStyle("Head2", parent=styles["Heading2"],
                                    fontSize=12, spaceBefore=12, spaceAfter=4,
                                    textColor=colors.HexColor("#1E3A5F"))
    body_style   = ParagraphStyle("Body2", parent=styles["Normal"],
                                   fontSize=9, spaceAfter=3, leading=13)
    small_style  = ParagraphStyle("Small", parent=styles["Normal"],
                                   fontSize=8, textColor=colors.grey)
    verdict_accepted = ParagraphStyle("VA", parent=body_style,
                                       textColor=colors.HexColor("#1a7a1a"))
    verdict_rejected = ParagraphStyle("VR", parent=body_style,
                                       textColor=colors.HexColor("#c0392b"))

    story = []

    # ── Header ────────────────────────────────────────────────────────────────
    story.append(Paragraph("⚖️ LITIGATION INTELLIGENCE OS", title_style))
    story.append(Paragraph("Evidence Checklist", heading_style))
    story.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor("#1E3A5F")))
    story.append(Spacer(1, 0.3*cm))

    # Case summary table
    from utils.helpers import parse_sections, format_currency
    sections = parse_sections(case.get("sections_violated", "[]"))
    summary_data = [
        ["Case Name", case.get("case_name", "—")],
        ["Client",    case.get("client_name", "—")],
        ["PAN",       case.get("assessee_pan", "—")],
        ["AY",        case.get("assessment_year", "—")],
        ["Sections",  ", ".join(sections)],
        ["Demand",    format_currency(case.get("demand_amount") or 0)],
        ["AO / Ward", f"{case.get('ao_name','—')} | {case.get('ao_ward','—')}"],
        ["Generated", datetime.now().strftime("%d %b %Y, %H:%M")],
    ]
    summary_table = Table(summary_data, colWidths=[3.5*cm, 12*cm])
    summary_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#EEF2F7")),
        ("FONTNAME",   (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTSIZE",   (0, 0), (-1, -1), 9),
        ("GRID",       (0, 0), (-1, -1), 0.3, colors.HexColor("#CBD5E0")),
        ("VALIGN",     (0, 0), (-1, -1), "TOP"),
        ("PADDING",    (0, 0), (-1, -1), 4),
    ]))
    story.append(summary_table)
    story.append(Spacer(1, 0.5*cm))

    # ── Evidence by section ───────────────────────────────────────────────────
    by_section: dict[str, list] = {}
    for item in evidence_items:
        sec = item.get("section", "General")
        by_section.setdefault(sec, []).append(item)

    for sec, items in by_section.items():
        story.append(Paragraph(f"Section {sec} — {len(items)} document(s)", heading_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#CBD5E0")))
        story.append(Spacer(1, 0.2*cm))

        for idx, item in enumerate(items, 1):
            verdict  = item.get("tribunal_verdict", "accepted")
            mandatory = item.get("is_mandatory") or item.get("mandatory", False)
            status    = item.get("status", "pending")
            count     = item.get("acceptance_count", 1)

            flag  = "🔴 MANDATORY" if mandatory else "🟡 Recommended"
            v_tag = "✅ ACCEPTED" if verdict == "accepted" else "❌ REJECTED"
            s_tag = {"available": "✓ Have it", "unavailable": "✗ Missing",
                     "pending": "⏳ Pending"}.get(status, status)

            # Document header row
            header_data = [[
                f"{idx}. {item.get('document_name','—')}",
                flag,
                v_tag,
                s_tag,
                f"Accepted in {count} case(s)",
            ]]
            header_tbl = Table(header_data,
                                colWidths=[6.5*cm, 2.8*cm, 2.5*cm, 2.2*cm, 3*cm])
            bg = colors.HexColor("#FFF3CD") if mandatory else colors.HexColor("#F8F9FA")
            header_tbl.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), bg),
                ("FONTNAME",   (0, 0), (0, 0), "Helvetica-Bold"),
                ("FONTSIZE",   (0, 0), (-1, -1), 8),
                ("GRID",       (0, 0), (-1, -1), 0.3, colors.HexColor("#DEE2E6")),
                ("PADDING",    (0, 0), (-1, -1), 4),
                ("VALIGN",     (0, 0), (-1, -1), "MIDDLE"),
            ]))
            story.append(header_tbl)

            # Why it matters / how to obtain
            why = item.get("why_it_matters", "")
            how = item.get("how_to_obtain", "")
            if why:
                story.append(Paragraph(f"<b>Why:</b> {why[:200]}", body_style))
            if how:
                story.append(Paragraph(f"<b>Obtain:</b> {how[:150]}", body_style))

            # Rejection reason in red if applicable
            rej = item.get("rejection_reason", "")
            if rej and verdict == "rejected":
                story.append(Paragraph(
                    f"<font color='#c0392b'><b>⚠ Rejection reason:</b> {rej[:200]}</font>",
                    body_style))

            story.append(Spacer(1, 0.15*cm))

        story.append(Spacer(1, 0.4*cm))

    # ── Footer summary ────────────────────────────────────────────────────────
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#1E3A5F")))
    mandatory_count = sum(1 for e in evidence_items if e.get("is_mandatory") or e.get("mandatory"))
    accepted_count  = sum(1 for e in evidence_items if e.get("tribunal_verdict") == "accepted")
    available_count = sum(1 for e in evidence_items if e.get("status") == "available")

    footer_data = [
        ["Total Documents", str(len(evidence_items)),
         "Mandatory", str(mandatory_count),
         "Tribunal-Accepted", str(accepted_count),
         "Have / Collected", str(available_count)],
    ]
    footer_tbl = Table(footer_data, colWidths=[3.5*cm, 1.5*cm, 2.5*cm, 1.5*cm, 3.5*cm, 1.5*cm, 3*cm, 1.5*cm])
    footer_tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#1E3A5F")),
        ("TEXTCOLOR",  (0, 0), (-1, -1), colors.white),
        ("FONTNAME",   (0, 0), (-1, -1), "Helvetica-Bold"),
        ("FONTSIZE",   (0, 0), (-1, -1), 9),
        ("ALIGN",      (0, 0), (-1, -1), "CENTER"),
        ("PADDING",    (0, 0), (-1, -1), 6),
    ]))
    story.append(Spacer(1, 0.3*cm))
    story.append(footer_tbl)
    story.append(Spacer(1, 0.3*cm))
    story.append(Paragraph(
        "Generated by Litigation Intelligence OS · Confidential · Not for public distribution",
        small_style))

    doc.build(story)
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# DOCX — Written Submissions
# ─────────────────────────────────────────────────────────────────────────────

def build_submission_docx(case: dict, submission_text: str,
                           title: str = "Written Submissions") -> bytes:
    """
    Build a properly formatted Word document for ITAT written submissions.
    Returns raw .docx bytes.
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml   import OxmlElement

    doc = Document()

    # Page setup — A4, 2.5cm margins
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = section.right_margin  = Cm(2.5)
    section.top_margin  = section.bottom_margin = Cm(2.5)

    # ── Styles ────────────────────────────────────────────────────────────────
    styles = doc.styles

    # Header style
    hdr_style = styles["Heading 1"]
    hdr_style.font.name  = "Arial"
    hdr_style.font.size  = Pt(14)
    hdr_style.font.bold  = True
    hdr_style.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    normal_style = styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(11)
    normal_style.paragraph_format.space_after  = Pt(6)
    normal_style.paragraph_format.line_spacing = Pt(18)

    from utils.helpers import parse_sections, format_currency
    sections = parse_sections(case.get("sections_violated", "[]"))

    # ── Title block ───────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("INCOME TAX APPELLATE TRIBUNAL")
    run.bold = True; run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    doc.add_paragraph()   # blank line

    case_title = doc.add_paragraph()
    case_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = case_title.add_run(case.get("case_name", "—"))
    run2.bold = True; run2.font.size = Pt(12)

    ay_para = doc.add_paragraph()
    ay_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ay_para.add_run(f"Assessment Year: {case.get('assessment_year','—')}  |  "
                    f"Demand: {format_currency(case.get('demand_amount') or 0)}")

    # Sections line
    sec_para = doc.add_paragraph()
    sec_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sec_run = sec_para.add_run(f"Sections: {', '.join(sections)}")
    sec_run.bold = True

    # Horizontal rule
    _add_hr(doc)

    doc.add_paragraph()

    # ── Submission content ────────────────────────────────────────────────────
    # Split by double newline → paragraphs; lines starting with # → headings
    paragraphs = submission_text.split("\n\n")
    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue

        if para_text.startswith("## "):
            p = doc.add_heading(para_text[3:].strip(), level=2)
            p.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
        elif para_text.startswith("# ") or para_text.startswith("**"):
            clean = para_text.lstrip("#").strip().strip("*").strip()
            p = doc.add_heading(clean, level=1)
            p.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
        elif para_text.startswith("- ") or para_text.startswith("• "):
            # Bullet list
            for line in para_text.split("\n"):
                line = line.lstrip("-• ").strip()
                if line:
                    p = doc.add_paragraph(line, style="List Bullet")
        else:
            # Inline bold handling
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            _add_formatted_run(p, para_text)

    # ── Footer ────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    _add_hr(doc)
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        f"Generated: {datetime.now().strftime('%d %b %Y')}  |  "
        "Litigation Intelligence OS  |  Confidential"
    )
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_paperbook_docx(case: dict, evidence_items: list[dict],
                          submission_text: str = "") -> bytes:
    """
    Build a paper book index DOCX — used as the cover sheet for documents
    filed before ITAT.
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = section.right_margin  = Cm(2.5)
    section.top_margin  = section.bottom_margin = Cm(2.5)

    from utils.helpers import parse_sections, format_currency
    sections = parse_sections(case.get("sections_violated", "[]"))

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("PAPER BOOK INDEX")
    r.bold = True; r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f"{case.get('case_name','—')}\n"
                 f"AY: {case.get('assessment_year','—')}  |  "
                 f"Sections: {', '.join(sections)}\n"
                 f"Demand: {format_currency(case.get('demand_amount') or 0)}")

    _add_hr(doc)
    doc.add_paragraph()

    # Table
    from docx.shared import Inches
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    headers = ["S.No", "Document Name", "Section", "Status", "Page No."]
    widths  = [Cm(1.2), Cm(8.5), Cm(2.5), Cm(2.8), Cm(2.0)]
    for i, (cell, hdr, w) in enumerate(zip(hdr_cells, headers, widths)):
        cell.text = hdr
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.width = w

    # Data rows
    for idx, item in enumerate(evidence_items, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = item.get("document_name", "—")[:60]
        row_cells[2].text = item.get("section", "—")
        row_cells[3].text = item.get("status", "pending").title()
        row_cells[4].text = "___"   # to be filled by advocate

        for cell in row_cells:
            cell.paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run("Note: Page numbers to be filled after assembling the paper book.")
    note.runs[0].font.size = Pt(9)
    note.runs[0].italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_full_case_pdf(case: dict, evidence_items: list[dict],
                         arguments: list[dict],
                         submission_text: str = "") -> bytes:
    """
    Build a single combined PDF:  Case Summary + Evidence Checklist + Arguments
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    Table, TableStyle, HRFlowable, PageBreak)

    buf    = io.BytesIO()
    doc    = SimpleDocTemplate(buf, pagesize=A4,
                                topMargin=2*cm, bottomMargin=2*cm,
                                leftMargin=2.5*cm, rightMargin=2.5*cm)
    styles = getSampleStyleSheet()

    title_style   = ParagraphStyle("T", parent=styles["Title"],  fontSize=18,
                                    textColor=colors.HexColor("#1E3A5F"), spaceAfter=6)
    heading_style = ParagraphStyle("H", parent=styles["Heading2"], fontSize=13,
                                    textColor=colors.HexColor("#1E3A5F"), spaceBefore=10)
    body_style    = ParagraphStyle("B", parent=styles["Normal"], fontSize=9, leading=14)
    small_style   = ParagraphStyle("S", parent=styles["Normal"], fontSize=8,
                                    textColor=colors.grey)

    story = []
    from utils.helpers import parse_sections, format_currency
    sections = parse_sections(case.get("sections_violated", "[]"))

    # Cover
    story.append(Spacer(1, 2*cm))
    story.append(Paragraph("⚖️ LITIGATION INTELLIGENCE OS", title_style))
    story.append(Paragraph("Full Case File", heading_style))
    story.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor("#1E3A5F")))
    story.append(Spacer(1, 0.5*cm))

    # Case info
    info_rows = [
        ["Case",     case.get("case_name","—")],
        ["Client",   case.get("client_name","—")],
        ["PAN",      case.get("assessee_pan","—")],
        ["AY",       case.get("assessment_year","—")],
        ["Sections", ", ".join(sections)],
        ["Demand",   format_currency(case.get("demand_amount") or 0)],
        ["Date",     datetime.now().strftime("%d %b %Y")],
    ]
    t = Table(info_rows, colWidths=[3*cm, 13*cm])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (0,-1), colors.HexColor("#EEF2F7")),
        ("FONTNAME",   (0,0), (0,-1), "Helvetica-Bold"),
        ("FONTSIZE",   (0,0), (-1,-1), 9),
        ("GRID",       (0,0), (-1,-1), 0.3, colors.HexColor("#CBD5E0")),
        ("PADDING",    (0,0), (-1,-1), 5),
    ]))
    story.append(t)
    story.append(PageBreak())

    # Evidence section
    story.append(Paragraph("Evidence Checklist", heading_style))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#CBD5E0")))
    story.append(Spacer(1, 0.3*cm))

    for idx, item in enumerate(evidence_items, 1):
        mandatory = item.get("is_mandatory") or item.get("mandatory", False)
        verdict   = item.get("tribunal_verdict", "accepted")
        bg = colors.HexColor("#FFF9E6") if mandatory else colors.HexColor("#F8F9FA")
        row = [[
            f"{idx}.",
            item.get("document_name","—"),
            "MANDATORY" if mandatory else "Optional",
            "✅" if verdict == "accepted" else "❌",
            item.get("status","—").title(),
        ]]
        rt = Table(row, colWidths=[0.7*cm, 8*cm, 2.2*cm, 1*cm, 2.1*cm])
        rt.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,-1), bg),
            ("FONTSIZE",   (0,0), (-1,-1), 8),
            ("GRID",       (0,0), (-1,-1), 0.2, colors.HexColor("#DEE2E6")),
            ("PADDING",    (0,0), (-1,-1), 3),
        ]))
        story.append(rt)
        why = item.get("why_it_matters","")
        if why:
            story.append(Paragraph(f"  → {why[:150]}", body_style))

    # Arguments section
    if arguments:
        story.append(PageBreak())
        story.append(Paragraph("Legal Arguments", heading_style))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#CBD5E0")))
        story.append(Spacer(1, 0.3*cm))
        for i, arg in enumerate(arguments[:10], 1):
            story.append(Paragraph(
                f"<b>Argument {i}:</b> {arg.get('argument_text','')[:400]}",
                body_style))
            if arg.get("source_citation"):
                story.append(Paragraph(
                    f"  Citation: {arg['source_citation'][:100]}",
                    small_style))
            story.append(Spacer(1, 0.2*cm))

    # Submission text
    if submission_text:
        story.append(PageBreak())
        story.append(Paragraph("Written Submissions", heading_style))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#CBD5E0")))
        story.append(Spacer(1, 0.3*cm))
        for line in submission_text.split("\n"):
            if line.strip():
                story.append(Paragraph(line.strip()[:400], body_style))

    story.append(Spacer(1, 0.5*cm))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#1E3A5F")))
    story.append(Paragraph(
        "Litigation Intelligence OS · Confidential · Not for public distribution",
        small_style))

    doc.build(story)
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def _add_hr(doc):
    from docx.oxml.ns import qn
    from docx.oxml   import OxmlElement
    from docx.shared import Pt, RGBColor
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"),  "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1E3A5F")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_formatted_run(para, text: str):
    """Handle **bold** inline markdown in a Word paragraph."""
    import re
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            para.add_run(part)
